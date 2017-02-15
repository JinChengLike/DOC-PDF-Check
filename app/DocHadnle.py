# coding=utf-8
import docx
import sys

reload(sys)
sys.setdefaultencoding("utf-8")


def readDocx(docName):
    fullText = []
    doc = docx.Document(docName)
    paras = doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    return fullText


def CheckDoc(mod, tes):
    len_1 = len(mod)
    len_2 = len(tes)
    codemessage = ""
    if len_1 != len_2:
        codemessage = "文件段落错误，请检查"
    i = 0
    while (i <= len_1):
        try:
            if mod[i] != tes[i]:
                return mod[i], tes[i]
            else:
                i += 1
        except IndexError:
            continue
        if i == len_1:
            break
    if "模版为" and "实际为" not in codemessage:
        codemessage = "文件无错误"
    return codemessage, codemessage

#
# if __name__ == '__main__':
#     doc = docx.Document('../instance/flies/jin.docx')
#     r = readDocx('../instance/flies/jin.docx')
#     print doc.paragraphs[1].text
#     # print doc.paragraphs[1].text
#     # print doc.paragraphs[2].text
#     # print doc.paragraphs[3].text
#     # print doc.paragraphs[4].text
